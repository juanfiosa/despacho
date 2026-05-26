"""
Generador de documentos Word (.docx) para el Poder Judicial de Córdoba.

Toma el texto ya renderizado por el motor Jinja2 y lo convierte en un .docx
con el formato estándar del expediente electrónico (Acuerdo Reglamentario
TSJA 1582).

Convenciones de formato:
- Fuente: Times New Roman 12pt (estándar en juzgados cordobeses)
- Márgenes: superior 2.5cm, inferior 2cm, izquierdo 3cm, derecho 2cm
- Interlineado: 1.5 líneas
- El archivo es editable para que el actuario pueda ajustar antes de firmar
"""

import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH


def _configurar_margenes(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)


def _estilo_parrafo(parrafo, negrita: bool = False, centrado: bool = False) -> None:
    parrafo.paragraph_format.space_before = Pt(0)
    parrafo.paragraph_format.space_after  = Pt(4)
    parrafo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if centrado:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()
    run.font.name    = "Times New Roman"
    run.font.size    = Pt(12)
    run.font.bold    = negrita
    run.font.color.rgb = RGBColor(0, 0, 0)


def _linea_separadora(doc: Document) -> None:
    p = doc.add_paragraph("_" * 48)
    _estilo_parrafo(p)


def texto_a_docx(texto_renderizado: str, nombre_archivo: str | None = None) -> bytes:
    """
    Convierte el texto renderizado por Jinja2 en un documento .docx.

    Args:
        texto_renderizado: string con el texto completo del decreto/auto
        nombre_archivo: se usa solo como referencia (no afecta el contenido)

    Returns:
        bytes del archivo .docx listo para guardar o enviar como respuesta HTTP
    """
    doc = Document()
    _configurar_margenes(doc)

    # Eliminar el párrafo vacío que crea python-docx por defecto
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    lineas = texto_renderizado.splitlines()

    for linea in lineas:
        contenido = linea.rstrip()

        # Línea de separación (firma)
        if contenido.startswith("___"):
            _linea_separadora(doc)
            continue

        # Título en mayúsculas (encabezados de sección)
        es_titulo = (
            contenido.isupper()
            and len(contenido) > 3
            and not contenido.startswith("$")
        )

        p = doc.add_paragraph()
        run = p.add_run(contenido)
        run.font.name  = "Times New Roman"
        run.font.size  = Pt(12)
        run.font.bold  = es_titulo
        run.font.color.rgb = RGBColor(0, 0, 0)

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2) if contenido else Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def guardar_docx(texto_renderizado: str, ruta: Path | str) -> None:
    """Guarda el .docx en disco."""
    contenido = texto_a_docx(texto_renderizado)
    Path(ruta).write_bytes(contenido)
